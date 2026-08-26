from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "generated"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

NAVY = "0B3B6E"
BLUE = "1479E8"
CYAN = "18A9E6"
INK = "132238"
MUTED = "5C6F86"
LINE = "D6E2F0"
PALE = "EEF6FF"
GREEN = "16875D"
AMBER = "C36A00"
RED = "C33434"
CODE_BG = "0D1B2A"
CODE_FG = "E8F1FA"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges):
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                edge.set(qn(f"w:{key}"), str(edge_data[key]))


def add_field(run, field_code):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def set_repeat_table_header(row):
    tr_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_properties.append(repeat)


def prevent_row_split(row):
    tr_properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_properties.append(cant_split)


def configure_document(doc, short_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 10),
        ("Subtitle", 12, MUTED, 0, 8),
        ("Heading 1", 17, NAVY, 15, 7),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 10.5, INK, 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(9.1)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.left_indent = Inches(0.22)
        style.paragraph_format.first_line_indent = Inches(-0.18)

    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(short_title.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(7.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Vendly.lk Engineering Guide  |  ")
    run.font.name = "Calibri"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    page_run = paragraph.add_run()
    add_field(page_run, "PAGE")
    page_run.font.name = "Calibri"
    page_run.font.size = Pt(7.5)
    page_run.font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc, eyebrow, title, subtitle, audience):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.75)
    set_cell_shading(cell, NAVY)
    set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
    cell.margin_top = Inches(0.35)
    cell.margin_bottom = Inches(0.35)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run(eyebrow.upper())
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(CYAN)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(title)
    r.font.name = "Calibri"
    r.font.size = Pt(27)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("D7E8FB")
    p = cell.add_paragraph()
    r = p.add_run(audience)
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("A9C9E8")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CURRENT PROJECT IMPLEMENTATION • AUGUST 2026")
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Workspace: D:\\Documents\\orderflow\\vendly-lk-web")
    r.font.name = "Consolas"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, text, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "6", "color": LINE},
        bottom={"val": "single", "sz": "6", "color": LINE},
        left={"val": "single", "sz": "18", "color": color},
        right={"val": "single", "sz": "6", "color": LINE},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(1)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, code, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(MUTED)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_BG)
    set_cell_border(cell, top={"val": "single", "sz": "5", "color": "294057"}, bottom={"val": "single", "sz": "5", "color": "294057"}, left={"val": "single", "sz": "5", "color": "294057"}, right={"val": "single", "sz": "5", "color": "294057"})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(code.strip())
    r.font.name = "Consolas"
    r.font.size = Pt(7.7)
    r.font.color.rgb = RGBColor.from_string(CODE_FG)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(value))
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(255, 255, 255)
        if widths:
            cell.width = Inches(widths[index])
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            if row_index % 2:
                set_cell_shading(cell, "F6F9FC")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            r.font.size = Pt(7.8)
            r.font.color.rgb = RGBColor.from_string(INK)
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": LINE})
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_contents(doc, sections):
    doc.add_heading("How to use this guide", level=1)
    add_para(doc, "Build the feature in order. Each section explains the reason first, shows a small code fragment, and then gives a verification step. Paths are relative to the Vendly project root unless an absolute path is shown.")
    add_table(doc, ["Part", "Purpose"], [[str(i + 1), title] for i, title in enumerate(sections)], [0.6, 5.9])
    add_callout(doc, "Learning rule", "Type the fragments yourself, run the verification after each stage, and commit only when that stage works. Do not paste the entire feature at once.")


def new_document(short_title):
    doc = Document()
    configure_document(doc, short_title)
    return doc


def build_chatbot_guide():
    doc = new_document("Customer Chatbot")
    add_cover(
        doc,
        "Vendly.lk implementation manual",
        "Customer Chatbot\nArchitecture & Step-by-Step Build Guide",
        "A source-aligned guide to the public storefront assistant, catalogue discovery, cart, order collection, reviews, voice input and Firestore persistence.",
        "Audience: a learner rebuilding the feature from scratch with React, Flask and Firestore",
    )
    sections = [
        "Architecture and request lifecycle",
        "Project files and prerequisites",
        "Public API and session security",
        "Firestore session and message model",
        "Conversation state machine",
        "Product discovery, details and alternatives",
        "Cart and order collection",
        "Order creation and post-order enquiries",
        "React storefront, polling and voice",
        "Reviews, media, testing and deployment",
    ]
    add_contents(doc, sections)

    doc.add_heading("1. Architecture and request lifecycle", level=1)
    add_para(doc, "The Vendly customer chatbot is a hybrid assistant. Deterministic Python rules own all transactional behavior; an optional AI provider only answers product questions and improves natural-language understanding. This separation protects stock, customer information and order creation from hallucinated actions.")
    add_code(doc, """
Customer browser (React StorefrontPage)
        |
        | HTTPS JSON + X-Chat-Session-Token
        v
Flask public blueprint (/api/v1/public)
        |
        +--> public_chat_service.py  [state machine + validation]
        |       +--> public_catalog_service.py
        |       +--> ai_service.py  [optional product answer]
        |       +--> review_service.py
        |       +--> order_service.py
        |
        v
Firestore
  publicChatSessions/{sessionId}
  publicChatSessions/{sessionId}/messages/{messageId}
  businesses/{businessId}/products, productVariants, orders, customers
""", "End-to-end architecture")
    add_table(doc, ["Layer", "Responsibility", "Trust boundary"], [
        ["React", "Displays messages/cards, keeps a temporary cart, polls for seller replies, captures voice.", "Never decides price, stock or final total."],
        ["Flask route", "Accepts HTTP input, applies rate limits and passes a verified session to services.", "Rejects malformed or unauthorized requests."],
        ["Chat service", "Owns states, parses required details, selects products and triggers an order.", "Only accepted state transitions are written."],
        ["Domain services", "Read products/reviews and create orders through shared business logic.", "Server recalculates totals and validates stock."],
        ["Firestore", "Stores session, messages, products, customer and order records.", "Admin SDK is backend-only."],
    ], [0.8, 3.0, 2.6])
    add_callout(doc, "Core invariant", "The browser may suggest a cart, but create_order() must re-read authoritative variants and compute totals on the server.", GREEN)

    doc.add_heading("2. Project files and prerequisites", level=1)
    add_table(doc, ["Current file", "Why it exists"], [
        ["frontend/src/pages/StorefrontPage.jsx", "Main storefront, chatbot UI, cart, checkout, reviews, voice and polling."],
        ["frontend/src/pages/StorefrontPage.css", "Responsive catalogue/chat/card/modal styling."],
        ["frontend/src/services/publicService.js", "Small frontend wrappers for public API calls."],
        ["frontend/src/services/apiClient.js", "Base URL, Firebase token, JSON parsing and retry behavior."],
        ["backend/app/api/public.py", "Public Flask endpoints and rate limits."],
        ["backend/app/services/public_chat_service.py", "Conversation state machine and secure session handling."],
        ["backend/app/services/ai_service.py", "Provider-neutral optional AI product answers."],
        ["backend/app/services/order_service.py", "Stock-safe order creation and totals."],
        ["backend/app/services/review_service.py", "Approved product and seller reviews."],
    ], [2.5, 4.0])
    doc.add_heading("2.1 Install the two applications", level=2)
    add_code(doc, """
# Backend
cd D:\\Documents\\orderflow\\vendly-lk-web\\backend
python -m venv .venv
.\\.venv\\Scripts\\Activate.ps1
pip install -r requirements.txt

# Frontend
cd D:\\Documents\\orderflow\\vendly-lk-web\\frontend
npm install
""", "PowerShell")
    add_heading = doc.add_heading
    add_heading("2.2 Configure environment variables", level=2)
    add_code(doc, """
# backend/.env
PORT=5000
FRONTEND_ORIGINS=http://localhost:5173,http://127.0.0.1:5173
FIREBASE_PROJECT_ID=your-project-id
FIREBASE_SERVICE_ACCOUNT_PATH=D:\\secrets\\firebase-admin.json
AI_PROVIDER=none
AI_API_KEY=
AI_MODEL=
RATE_LIMIT_STORAGE_URI=memory://

# frontend/.env
VITE_FIREBASE_API_KEY=...
VITE_FIREBASE_AUTH_DOMAIN=...
VITE_FIREBASE_PROJECT_ID=...
VITE_FIREBASE_APP_ID=...
VITE_API_BASE_URL=http://127.0.0.1:5000/api/v1
""")
    add_callout(doc, "Do not commit secrets", "Commit .env.example files, but keep .env and the Firebase service-account JSON in .gitignore.", RED)

    doc.add_heading("3. Public API and session security", level=1)
    add_table(doc, ["Method", "Endpoint", "Purpose", "Rate"], [
        ["GET", "/api/v1/public/stores/{code}", "Load seller storefront and catalogue.", "Public"],
        ["POST", "/api/v1/public/chat/sessions", "Create a token-protected conversation.", "Public"],
        ["POST", "/api/v1/public/chat/sessions/{id}/messages", "Send one customer message and receive assistant output.", "60/min"],
        ["GET", "/api/v1/public/chat/sessions/{id}/messages", "Load/poll persisted history.", "120/min"],
        ["POST", "/api/v1/public/chat/sessions/{id}/orders", "Create an order from the collected cart/details.", "6/min"],
        ["POST", "/api/v1/public/chat/sessions/{id}/claim", "Attach a guest chat to a signed-in customer.", "Authenticated"],
        ["GET", "/api/v1/public/products/{code}/reviews", "Load approved public reviews.", "Public"],
    ], [0.55, 2.9, 2.55, 0.65])
    add_code(doc, """
@public_blueprint.post("/chat/sessions/<session_id>/messages")
@limiter.limit("60 per minute", key_func=public_chat_key)
def public_chat_message(session_id):
    result = answer_public_message(
        get_firestore_client(),
        session_id,
        request.headers.get("X-Chat-Session-Token", ""),
        get_json_object(),
    )
    return jsonify({"assistant": result})
""", "backend/app/api/public.py")
    add_para(doc, "The route is intentionally thin. Business logic belongs in the service so it can be tested without starting Flask.")
    add_heading("3.1 Store only a hash of the chat token", level=2)
    add_code(doc, """
def token_hash(token):
    return hashlib.sha256(token.encode("utf-8")).hexdigest()

token = secrets.token_urlsafe(32)
session_reference.set({
    "tokenHash": token_hash(token),
    "status": "active",
    "state": "browsing",
    "cart": [],
})
return {"id": session_reference.id, "sessionToken": token}
""")
    add_para(doc, "The raw token is returned once to the browser and kept in sessionStorage/localStorage. Later requests hash the supplied token and compare it with hmac.compare_digest(), reducing exposure if Firestore data is leaked.")

    doc.add_heading("4. Firestore session and message model", level=1)
    add_table(doc, ["Path / field", "Type", "Meaning"], [
        ["publicChatSessions/{id}", "document", "One conversation owned by a store/product short link."],
        ["businessId / productId", "string", "Tenant scope and optional product-link restriction."],
        ["tokenHash", "string", "SHA-256 hash of the browser session token."],
        ["customerUid", "string", "Firebase UID after guest session is claimed."],
        ["state", "string", "Current deterministic conversation step."],
        ["cart", "array", "Only variantId and positive quantity."],
        ["customerDraft", "map", "Name, phones, address, district, city and note while collecting."],
        ["orderId", "string", "Linked completed order, kept for status questions."],
        ["messages/{messageId}", "subcollection", "Role, text, metadata and createdAt timestamp."],
    ], [2.3, 0.75, 3.6])
    add_code(doc, """
def save_chat_message(session_reference, role, message, metadata=None):
    session_reference.collection("messages").document().set({
        "role": role,
        "message": message,
        "metadata": metadata or {},
        "createdAt": firestore.SERVER_TIMESTAMP,
    })
""")
    add_callout(doc, "Why a subcollection?", "Messages can grow independently without making the session document exceed Firestore's document size limit.")

    doc.add_heading("5. Conversation state machine", level=1)
    add_code(doc, """
browsing
  -> collecting-name
  -> collecting-phone
  -> collecting-secondary-phone
  -> collecting-address
  -> collecting-district
  -> collecting-nearest-city
  -> collecting-delivery-note
  -> awaiting-confirmation
  -> completed

completed -- explicit "another order" --> browsing
completed -- status question ----------> completed (same order context)
""", "Allowed state transitions")
    add_table(doc, ["State", "Accepted input", "Next state"], [
        ["browsing", "Product/category questions, Add buttons, or explicit order intent.", "Browsing or collecting-name"],
        ["collecting-name", "A plausible customer name.", "collecting-phone"],
        ["collecting-phone", "Valid Sri Lankan primary number.", "collecting-secondary-phone"],
        ["collecting-secondary-phone", "Valid second number or skip phrase.", "collecting-address"],
        ["collecting-address", "Street/delivery address.", "collecting-district"],
        ["collecting-district", "District.", "collecting-nearest-city"],
        ["collecting-nearest-city", "Nearest city/town.", "collecting-delivery-note"],
        ["collecting-delivery-note", "Note or none.", "awaiting-confirmation"],
        ["awaiting-confirmation", "Confirm or correction request.", "completed or earlier state"],
        ["completed", "Order status/info or explicit new order.", "completed or browsing"],
    ], [1.45, 3.25, 1.75])
    add_code(doc, """
def respond(text, action, *, next_state=None, metadata=None):
    state = next_state or current_state
    save_chat_message(session_reference, "assistant", text, {
        "action": action,
        "state": state,
        **(metadata or {}),
    })
    session_reference.update({
        "state": state,
        "cart": cart,
        "customerDraft": customer_draft,
        "updatedAt": firestore.SERVER_TIMESTAMP,
    })
    return {"message": text, "action": action, "state": state,
            "cartSummary": cart_summary}
""", "A single response helper keeps persistence consistent")

    doc.add_heading("6. Product discovery, details and alternatives", level=1)
    add_steps(doc, [
        "Load only active products belonging to the session business.",
        "Try exact/product-token matching with find_product_in_message().",
        "If the message is a category request, return all products in that category.",
        "For one specific product, return its photos, description and approved reviews.",
        "Offer related products from the same category only after the customer is not satisfied.",
        "Do not add anything to cart until the customer clicks Add or clearly asks to order it.",
    ])
    add_code(doc, """
def related_products(products, selected_product, limit=4):
    return [
        product for product in products
        if product["id"] != selected_product["id"]
        and product.get("categoryId") == selected_product.get("categoryId")
    ][:limit]
""")
    add_heading("6.1 Optional AI answers must be fact-bounded", level=2)
    add_code(doc, """
def product_prompt(question, product):
    facts = {
        "name": product.get("name"),
        "brand": product.get("brand"),
        "description": product.get("description"),
        "sellerAiDescription": product.get("aiDescription"),
        "priceLkr": product.get("sellingPriceMinor", 0) / 100,
        "availableSizes": [v.get("size") for v in product.get("variants", [])],
        "approvedReviewSnippets": product.get("approvedReviewSnippets", []),
    }
    return f"Use only PRODUCT FACTS. Keep the answer brief.\\nQUESTION: {question}\\nPRODUCT FACTS: {facts}"
""")
    add_callout(doc, "No AI web browsing in the order transaction", "If a specification is missing, say it is not confirmed or use a separately governed product-enrichment workflow. Never invent waterproofing, SIM, warranty or health claims.", AMBER)

    doc.add_heading("7. Cart and order collection", level=1)
    add_heading("7.1 Normalize browser cart input", level=2)
    add_code(doc, """
def normalize_chat_cart(value):
    result = []
    for item in value if isinstance(value, list) else []:
        variant_id = str(item.get("variantId") or "").strip()
        try:
            quantity = int(item.get("quantity") or 0)
        except (TypeError, ValueError):
            continue
        if variant_id and 0 < quantity <= 1000:
            result.append({"variantId": variant_id, "quantity": quantity})
    return result
""")
    add_para(doc, "The server then looks up every variant and removes unknown IDs. It returns cartSummary so the UI displays server-approved names, images, prices, quantities and line totals.")
    add_heading("7.2 Begin detail collection only after selection is finished", level=2)
    add_code(doc, """
if current_state == "browsing" and cart_summary \
        and is_finished_selecting_items(message):
    return respond(
        f"Great. Your draft contains {sum(i['quantity'] for i in cart_summary)} item(s). What is your name?",
        "collect-name",
        next_state="collecting-name",
    )
""")
    add_heading("7.3 Validate customer details one or two at a time", level=2)
    add_code(doc, """
if current_state == "collecting-phone":
    try:
        phone = normalize_sri_lankan_phone(message)
    except ValueError as error:
        return respond(str(error), "collect-phone", next_state="collecting-phone")
    customer_draft["phoneNumber"] = phone
    return respond(
        "Do you have a second phone number? Send it, or type skip.",
        "collect-secondary-phone",
        next_state="collecting-secondary-phone",
    )
""")
    add_callout(doc, "User experience rule", "Ask for missing information progressively. The required order data is product, quantity, name, primary phone, address, district and nearest city; the second phone and note are optional.")

    doc.add_heading("8. Order creation and post-order enquiries", level=1)
    add_heading("8.1 Confirm before writing", level=2)
    add_code(doc, """
if current_state == "awaiting-confirmation":
    if normalized_phrase(message) in CONFIRMATION_PHRASES:
        order = create_public_chat_order(database, session_id, provided_token, {
            "items": cart,
            "customer": customer_draft,
            "deliveryNote": customer_draft.get("deliveryNote", ""),
        })
        # Persist order link, clear cart, retain customer summary.
        return respond(
            f"Your order {order['orderNumber']} was placed successfully.",
            "order-confirmed",
            next_state="completed",
        )
""")
    add_para(doc, "create_public_chat_order() finds or creates the customer, builds an order payload, enforces product-link restrictions, calls the shared create_order() service, and saves orderId on the chat session.")
    add_heading("8.2 Keep a completed chat open for support", level=2)
    add_code(doc, """
if current_state == "completed" or session.get("status") == "completed":
    if is_explicit_new_order_request(message):
        # Clear only the cart and start a new order flow.
        return respond("What product would you like to order?", "prompt-product",
                       next_state="browsing")
    if is_order_enquiry(message):
        order = latest_order_for_session(database, session)
        return respond(order_information_message(order), "show-order-info",
                       next_state="completed")
    return respond("Your order is saved. Ask me for its status or say 'another order'.",
                   "order-help", next_state="completed")
""")
    add_callout(doc, "Bug prevention", "Do not mark the conversation closed after checkout. Keep status=completed accepted by authorization; only an explicit closed/disabled state should block messaging.", GREEN)

    doc.add_heading("9. React storefront, polling and voice", level=1)
    add_heading("9.1 Keep API wrappers small", level=2)
    add_code(doc, """
export function sendPublicChatMessage(sessionId, sessionToken, data) {
  return apiRequest(`/public/chat/sessions/${sessionId}/messages`, {
    method: "POST",
    body: data,
    headers: { "X-Chat-Session-Token": sessionToken },
    requiresAuthentication: false,
  });
}
""", "frontend/src/services/publicService.js")
    add_heading("9.2 Send cart with every chat request", level=2)
    add_code(doc, """
const response = await sendPublicChatMessage(session.id, session.sessionToken, {
  message: cleanMessage,
  cart: cart.map((item) => ({
    variantId: item.variantId,
    quantity: item.quantity,
  })),
});

setSession((current) => ({ ...current, state: response.assistant.state }));
setMessages((current) => [...current, {
  role: "assistant",
  text: response.assistant.message,
  action: response.assistant.action,
  products: response.assistant.products,
}]);
""")
    add_heading("9.3 Poll for seller and status messages", level=2)
    add_code(doc, """
useEffect(() => {
  if (!session) return undefined;
  const timer = window.setInterval(async () => {
    const result = await getPublicChatMessages(session.id, session.sessionToken);
    appendOnlyUnseenMessages(result.messages);
  }, 4000);
  return () => window.clearInterval(timer);
}, [session]);
""")
    add_para(doc, "The current implementation tracks received message IDs in a ref, preventing duplicates. For larger scale, replace polling with a controlled realtime channel or Firestore listener exposed through secure rules.")
    add_heading("9.4 Add voice input as progressive enhancement", level=2)
    add_code(doc, """
const SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;
if (!SpeechRecognition) {
  setErrorMessage("Voice input is not supported in this browser.");
  return;
}
const recognition = new SpeechRecognition();
recognition.lang = voiceLanguage;      // en-LK or si-LK
recognition.interimResults = false;
recognition.onresult = (event) => {
  const transcript = event.results?.[0]?.[0]?.transcript?.trim();
  if (transcript) sendChatText(transcript);
};
recognition.start();
""")
    add_callout(doc, "Browser requirement", "Camera/microphone features normally require HTTPS or localhost. A phone opening a plain LAN IP may block them even when the code is correct.", AMBER)

    doc.add_heading("10. Reviews, media, testing and deployment", level=1)
    add_heading("10.1 Return only approved reviews", level=2)
    add_code(doc, """
reviews = list_public_product_reviews(database, product["id"])
seller_rating = list_public_seller_reviews(database, session["businessId"])
return respond(
    "Here are verified reviews for this product and the seller rating.",
    "show-reviews",
    metadata={"reviews": reviews, "sellerRating": seller_rating},
)
""")
    add_para(doc, "Review images should be stored as Cloudinary HTTPS URLs. Render them as links with target=\"_blank\" and rel=\"noreferrer\" so the full image opens safely in a new tab.")
    add_heading("10.2 Minimum automated tests", level=2)
    add_table(doc, ["Test", "Expected result"], [
        ["Create session with valid store code", "Returns sessionId + raw token; Firestore stores only tokenHash."],
        ["Wrong session token", "401/403; no message or order is written."],
        ["Unknown variant in cart", "Removed or rejected; never affects stock."],
        ["Invalid phone/address", "State remains at the same collection step."],
        ["Multiple products", "One order with many items, not separate orders."],
        ["Confirm order", "Order created once; stock reserved once; totals server-computed."],
        ["Ask status after checkout", "Returns linked order status; does not reset catalogue."],
        ["Say another order", "Starts a clean browsing flow while keeping history."],
        ["AI provider unavailable", "Deterministic catalogue/order flow still works."],
    ], [2.5, 4.0])
    add_heading("10.3 Run and verify", level=2)
    add_code(doc, """
# Terminal 1
cd backend
.\\.venv\\Scripts\\Activate.ps1
python run.py

# Terminal 2
cd frontend
npm run dev -- --host 0.0.0.0

# Backend tests
cd backend
pytest -q

# Production frontend check
cd frontend
npm run build
""")
    add_heading("10.4 Troubleshooting checklist", level=2)
    add_table(doc, ["Symptom", "Check"], [
        ["Failed to fetch", "Backend running, VITE_API_BASE_URL, CORS origin, Windows firewall and same LAN."],
        ["Chat session is closed", "Authorization accepts active and completed; checkout did not set closed."],
        ["Cards appear but cart is empty", "Add handler updates cart and sends variantId/quantity on next message."],
        ["Images are cropped", "Use object-fit: contain; fixed media area; preserve source aspect ratio."],
        ["Messages require refresh", "Polling interval active and unseen IDs are merged, not overwritten."],
        ["AI replies incorrectly", "Keep transaction states deterministic and reduce prompt scope to product facts."],
        ["Quota errors", "Use Firestore emulator locally and avoid polling too frequently."],
    ], [2.2, 4.3])
    add_callout(doc, "Completion definition", "The chatbot is complete only when it can discover products, answer fact-bounded questions, collect a multi-item order, confirm once, persist history, show later status updates, and continue support after checkout.", GREEN)
    path = OUTPUT_DIR / "Vendly_Customer_Chatbot_Build_Guide.docx"
    doc.save(path)
    return path


def build_business_assistant_guide():
    doc = new_document("Business Assistant")
    add_heading = doc.add_heading
    add_cover(
        doc,
        "Vendly.lk implementation manual",
        "Seller Business Assistant\nArchitecture & Step-by-Step Build Guide",
        "A source-aligned guide to permission-aware dashboard search, navigation, analytics, exports, confirmed writes, audit logging and voice interaction.",
        "Audience: a learner rebuilding the seller assistant from scratch with React, Flask and Firestore",
    )
    sections = [
        "Architecture and safety model",
        "Files, authentication and permissions",
        "Assistant endpoint and response contract",
        "Deterministic intents and AI fallback",
        "Read-only business operations",
        "Client actions and dashboard navigation",
        "Confirmed write operations and audit logs",
        "React panel, cards and voice",
        "Testing, extension and production readiness",
    ]
    add_contents(doc, sections)

    doc.add_heading("1. Architecture and safety model", level=1)
    add_para(doc, "The Business Assistant is not a general autonomous agent. It is a permission-aware command interpreter placed above existing Vendly services. It may summarize and navigate immediately, but any stock or order mutation is prepared first and executed only after explicit seller confirmation.")
    add_code(doc, """
Seller dashboard (React BusinessAssistant)
        |
        | Firebase bearer token + businessId + message
        v
Flask endpoint /businesses/{businessId}/assistant/messages
        |
        +--> require_firebase_user
        +--> require_business_member
        +--> deterministic_intent(message)
        |       \\--> optional AI classification fallback
        |
        +--> process_read_intent()  ------> immediate response/clientAction
        \\--> prepare_write_intent() ----> confirmation card
                    |
                    | seller confirms
                    v
              confirm_action() --> domain service --> Firestore + audit log
""", "Safety-first request lifecycle")
    add_table(doc, ["Rule", "Implementation"], [
        ["Tenant isolation", "Every request includes businessId and verifies an active membership document."],
        ["Least privilege", "Each intent calls require_permission() for orders, products, customers or analytics."],
        ["No direct AI writes", "AI returns a structured intent; server resolves real IDs and allowed values."],
        ["Human confirmation", "Mutations return pendingAction and require a second request with confirmedAction."],
        ["Allowlist", "Only update_order_status and adjust_stock are executable confirmed actions."],
        ["Auditability", "Confirmed actions are copied into assistantAuditLogs with actor UID and result IDs."],
    ], [1.45, 5.05])

    doc.add_heading("2. Files, authentication and permissions", level=1)
    add_table(doc, ["Current file", "Responsibility"], [
        ["frontend/src/components/BusinessAssistant.jsx", "Floating panel, messages, suggestions, cards, confirmations, exports and voice."],
        ["frontend/src/components/BusinessAssistant.css", "Panel, mobile layout, message and action styling."],
        ["frontend/src/services/businessAssistantService.js", "POST wrapper for message or confirmedAction."],
        ["backend/app/api/business_assistant.py", "Authenticated, rate-limited Flask endpoint."],
        ["backend/app/services/business_assistant_service.py", "Intent classification, permission checks, reads, prepared writes and audits."],
        ["backend/app/services/ai_service.py", "Optional structured intent classification."],
        ["backend/app/core/auth.py", "Firebase ID-token verification and verified-email rule."],
        ["backend/app/core/authorization.py", "Membership and resource permission enforcement."],
    ], [2.75, 3.75])
    add_heading("2.1 Protect the endpoint", level=2)
    add_code(doc, """
@business_assistant_blueprint.post(
    "/businesses/<business_id>/assistant/messages"
)
@limiter.limit("30 per minute")
@require_firebase_user
@require_business_member()
def business_assistant_message(business_id):
    response = handle_business_assistant_message(
        get_firestore_client(), business_id,
        g.current_user["uid"], g.membership, get_json_object(),
    )
    return jsonify({"assistant": response})
""", "backend/app/api/business_assistant.py")
    add_para(doc, "Decorator order matters: requests are rate-limited, authenticated and authorized before the service receives a database client or membership.")
    add_heading("2.2 Verify fine-grained permission inside each intent", level=2)
    add_code(doc, """
def require_permission(membership, permission):
    if not membership_has_permission(membership, permission):
        raise ApiError(
            "permission_denied",
            "Your staff role does not allow this assistant action.",
            403,
            {"requiredPermission": permission},
        )
""")
    add_table(doc, ["Intent family", "Suggested permission"], [
        ["Order search/filter/status", "orders:read / orders:update"],
        ["Inventory search/adjust", "products:read / products:update"],
        ["Customer search/export", "customers:read / customers:export"],
        ["Business summary", "analytics:read"],
        ["Settings/staff", "settings:read or members:manage"],
    ], [2.4, 4.1])

    doc.add_heading("3. Assistant endpoint and response contract", level=1)
    add_code(doc, """
// Request: natural-language message
{ "message": "Show packed orders" }

// Response
{
  "assistant": {
    "message": "Opening packed orders.",
    "cards": [],
    "suggestions": ["Reset filters"],
    "navigation": "/orders?status=packed",
    "clientAction": null,
    "pendingAction": null
  }
}

// Request: confirm a prepared mutation
{ "confirmedAction": { "type": "adjust_stock", ... } }
""", "Stable JSON contract")
    add_table(doc, ["Field", "Meaning"], [
        ["message", "Text shown in the assistant bubble."],
        ["cards", "Small order/product/business summary objects for visual rendering."],
        ["suggestions", "Safe next prompts displayed as buttons."],
        ["navigation", "Internal React route the browser may open."],
        ["clientAction", "Browser-only action such as export, modal open, settings panel or reset filters."],
        ["pendingAction", "Server-prepared mutation requiring explicit confirmation."],
    ], [1.35, 5.15])
    add_code(doc, """
export async function sendBusinessAssistantMessage(
  businessId,
  { message = "", confirmedAction = null } = {},
) {
  const response = await apiRequest(
    `/businesses/${businessId}/assistant/messages`,
    { method: "POST", body: confirmedAction ? { confirmedAction } : { message } },
  );
  return response.assistant;
}
""", "frontend/src/services/businessAssistantService.js")

    doc.add_heading("4. Deterministic intents and AI fallback", level=1)
    add_para(doc, "Fast, common commands are classified with Python first. Only unknown language is sent to the optional provider. This keeps the assistant fast, inexpensive and testable.")
    add_code(doc, """
def classify_intent(message):
    deterministic = deterministic_intent(message)
    if deterministic["intent"] != "unknown":
        return deterministic
    return generate_business_assistant_intent(message) or deterministic
""")
    add_heading("4.1 Build exact deterministic patterns", level=2)
    add_code(doc, """
def deterministic_intent(message):
    lowered = message.casefold().strip()
    if lowered in {"help", "what can you do"}:
        return {"intent": "help"}
    if "export orders" in lowered:
        return {"intent": "export_orders"}
    if "add product" in lowered:
        return {"intent": "open_add_product"}
    if "low stock" in lowered:
        return {"intent": "low_stock"}
    match = ORDER_NUMBER_PATTERN.search(message)
    if match:
        return {"intent": "search_order", "orderQuery": match.group(0)}
    return {"intent": "unknown"}
""")
    add_heading("4.2 Force AI to return structured JSON", level=2)
    add_code(doc, """
ALLOWED_INTENTS = {
  "business_summary", "search_order", "search_product", "order_view",
  "inventory_view", "update_order_status", "adjust_stock", "navigate"
}

prompt = (
  "Return one JSON object only. Use an allowed intent. "
  "Never invent an ID, status or quantity not present in SELLER MESSAGE. "
  "For a stock change return productQuery, optional variantQuery and signed "
  "quantityChange. For an order update return orderQuery and status."
)
""")
    add_callout(doc, "Validate after AI", "AI output is untrusted. Recheck intent, status, quantity, record match and membership before acting.", RED)

    doc.add_heading("5. Read-only business operations", level=1)
    add_table(doc, ["Intent", "Server behavior", "Typical output"], [
        ["business_summary", "get_business_analytics()", "Revenue, order and stock summary cards."],
        ["pending_orders", "list_orders(status='needs-confirmation')", "Count and order cards."],
        ["low_stock", "list_products() then filter stockStatus", "Product cards and inventory link."],
        ["search_order", "Find exact order number or matching customer.", "One or more order cards."],
        ["search_product", "Token-match product name/SKU/barcode.", "Product cards."],
        ["order_view", "Build /orders query parameters.", "Internal navigation."],
        ["inventory_view", "Build /inventory query parameters and sort.", "Internal navigation."],
        ["customer_view", "Open customer page with query.", "Internal navigation."],
    ], [1.25, 3.4, 1.85])
    add_code(doc, """
if intent_name == "business_summary":
    require_permission(membership, "analytics:read")
    analytics = get_business_analytics(database, business_id)
    return {
        "message": "Here is today's business summary.",
        "cards": [{
            "type": "summary",
            "orders": analytics["orders"]["all"],
            "revenue": format_money(analytics["revenueMinor"]),
            "lowStock": analytics["inventory"]["lowStock"],
        }],
    }
""")
    add_para(doc, "Reuse existing services instead of duplicating Firestore queries inside the assistant. The assistant becomes an orchestration layer, while analytics/order/product services remain the single source of truth.")

    doc.add_heading("6. Client actions and dashboard navigation", level=1)
    add_table(doc, ["clientAction.type", "React behavior"], [
        ["export_orders", "Call downloadOrderExport() and save the returned file."],
        ["export_inventory", "Load products and call downloadInventoryCsv()."],
        ["export_customers", "Load customers and call downloadCustomersCsv()."],
        ["open_settings", "Dispatch a settings event with the requested section."],
        ["reset_filters", "Dispatch vendly:reset-filters for the active page."],
        ["open_add_order/product/courier", "Navigate and dispatch an event that opens the corresponding modal."],
    ], [2.1, 4.4])
    add_code(doc, """
async function executeResponseAction(response) {
  const action = response?.clientAction;
  if (action?.type === "export_orders") {
    await downloadOrderExport(business.id);
  }
  if (action?.type === "open_settings") {
    window.dispatchEvent(new CustomEvent("vendly:open-settings", {
      detail: { section: action.section || "general" },
    }));
  }
  if (response?.navigation) navigate(response.navigation);
}
""", "frontend/src/components/BusinessAssistant.jsx")
    add_callout(doc, "Keep UI work in the browser", "The backend describes the action; React owns navigation, modal visibility and file downloads. This avoids coupling Flask to browser implementation details.")

    # Keep the confirmation workflow together; splitting its first code sample
    # across pages makes the safety sequence harder for a learner to follow.
    doc.add_page_break()
    doc.add_heading("7. Confirmed write operations and audit logs", level=1)
    add_heading("7.1 Prepare, do not execute", level=2)
    add_code(doc, """
if intent_name == "update_order_status":
    require_permission(membership, "orders:update")
    orders = find_order(database, business_id, intent.get("orderQuery"))
    if len(orders) != 1:
        return {"message": "I need one exact order number."}
    order = orders[0]
    return {
        "message": f"Please confirm: change {order['orderNumber']} to {status}?",
        "pendingAction": {
            "type": "update_order_status",
            "orderId": order["id"],
            "status": status,
            "label": f"Mark {order['orderNumber']} as {status}",
        },
    }
""")
    add_heading("7.2 Confirm through an allowlist", level=2)
    add_code(doc, """
ALLOWED_CONFIRMED_ACTIONS = {"update_order_status", "adjust_stock"}

def confirm_action(database, business_id, uid, membership, action):
    if action.get("type") not in ALLOWED_CONFIRMED_ACTIONS:
        raise ApiError("invalid_assistant_action", "Action is not allowed.", 422)
    if action["type"] == "update_order_status":
        require_permission(membership, "orders:update")
        order = update_order_status(
            database, business_id, action["orderId"], uid,
            {"status": action["status"],
             "note": "Updated through the Vendly business assistant"},
        )
        audit_action(database, business_id, uid, action, order)
        return {"message": f"Done — {order['orderNumber']} is now {action['status']}."}
""")
    add_heading("7.3 Save an audit event", level=2)
    add_code(doc, """
def audit_action(database, business_id, uid, action, result):
    database.collection("businesses").document(business_id) \
        .collection("assistantAuditLogs").document().set({
            "actorUid": uid,
            "action": action,
            "resultId": result.get("id"),
            "createdAt": firestore.SERVER_TIMESTAMP,
        })
""")
    add_callout(doc, "Concurrency", "Stock changes must use the existing stock service/transaction. A confirmation card is not a stock lock; revalidate available/reserved stock when the confirmed request executes.", AMBER)

    doc.add_page_break()
    doc.add_heading("8. React panel, cards and voice", level=1)
    add_heading("8.1 Message flow", level=2)
    add_code(doc, """
async function sendMessage(text) {
  const cleanMessage = text.trim();
  if (!cleanMessage || !business?.id || isSending) return;
  setMessages((items) => [...items, {
    id: messageId("seller"), role: "user", text: cleanMessage,
  }]);
  setIsSending(true);
  try {
    const response = await sendBusinessAssistantMessage(business.id, {
      message: cleanMessage,
    });
    appendAssistantResponse(response);
    await executeResponseAction(response);
  } finally {
    setIsSending(false);
  }
}
""")
    add_heading("8.2 Confirm or cancel a prepared action", level=2)
    add_code(doc, """
async function confirmAction(action, sourceMessageId) {
  markActionState(sourceMessageId, "confirmed");
  const response = await sendBusinessAssistantMessage(business.id, {
    confirmedAction: action,
  });
  appendAssistantResponse(response);
}

function cancelAction(sourceMessageId) {
  markActionState(sourceMessageId, "cancelled");
}
""")
    add_heading("8.3 Voice input and spoken replies", level=2)
    add_code(doc, """
const recognition = new SpeechRecognition();
recognition.lang = voiceLanguage; // en-LK or si-LK
recognition.interimResults = false;
recognition.onresult = (event) => {
  const transcript = event.results?.[0]?.[0]?.transcript?.trim() || "";
  if (transcript) sendMessage(transcript);
};

function speak(text) {
  if (!speechEnabled || !window.speechSynthesis || !text) return;
  const utterance = new SpeechSynthesisUtterance(text);
  utterance.lang = voiceLanguage;
  window.speechSynthesis.cancel();
  window.speechSynthesis.speak(utterance);
}
""")
    add_bullets(doc, [
        "Always provide text input; voice is optional progressive enhancement.",
        "Show listening state and a stop control.",
        "Cancel speech synthesis when the panel closes or a new reply starts.",
        "Never execute a mutation directly from a transcript; confirmation still applies.",
    ])

    doc.add_page_break()
    doc.add_heading("9. Testing, extension and production readiness", level=1)
    add_heading("9.1 Minimum tests", level=2)
    add_table(doc, ["Test", "Expected result"], [
        ["Unauthenticated request", "401 before any assistant work."],
        ["User outside business", "403 tenant access denied."],
        ["Staff missing permission", "403 with requiredPermission detail."],
        ["Deterministic export/navigation", "Correct clientAction without AI call."],
        ["AI returns invalid JSON", "Falls back to help/unknown safely."],
        ["Ambiguous order/product", "Asks for one exact match; no write."],
        ["Write request", "Returns pendingAction; database unchanged."],
        ["Tampered confirmedAction", "Rejected by type/status/quantity validation."],
        ["Valid confirmation", "Shared service changes data and audit log is written."],
        ["Concurrent stock change", "Domain transaction prevents invalid stock."],
    ], [2.45, 4.05])
    add_heading("9.2 Add a new intent safely", level=2)
    add_steps(doc, [
        "Name the intent and decide whether it is read-only, browser-only or mutating.",
        "Add deterministic keywords for common phrasing.",
        "If AI may classify it, add the name and exact JSON fields to the prompt allowlist.",
        "Add permission enforcement before reading or preparing data.",
        "For a browser action, return clientAction and implement it in React.",
        "For a mutation, prepare a minimal pendingAction containing resolved server IDs.",
        "Add the type to ALLOWED_CONFIRMED_ACTIONS only after confirm_action validates it.",
        "Call an existing domain service and write an audit record.",
        "Add unit tests for permission denied, ambiguous target, tampering and success.",
    ])
    add_heading("9.3 Recommended next capabilities", level=2)
    add_table(doc, ["Capability", "Safe first version"], [
        ["Courier recommendation", "Read-only ranked explanation; seller selects courier."],
        ["Dead-stock report", "Read product age and sales history; navigate to filtered inventory."],
        ["Profit question", "Read-only analytics cards using verified cost and sales data."],
        ["Daily work centre", "Return actionable counts and links, not autonomous writes."],
        ["Create promotion", "Prepare a draft campaign; require confirmation before publish."],
        ["Role management", "Open settings only; do not let natural language change owners initially."],
    ], [2.0, 4.5])
    add_heading("9.4 Production checklist", level=2)
    add_bullets(doc, [
        "Use HTTPS and exact CORS origins; never wildcard credentialed production requests.",
        "Use Redis-backed shared rate-limit storage when running multiple backend instances.",
        "Keep AI keys and Firebase Admin credentials on Flask only.",
        "Record provider latency/errors without logging seller messages containing personal data.",
        "Apply timeouts and retain deterministic behavior when the AI provider fails.",
        "Monitor assistantAuditLogs and alert on repeated rejected/tampered actions.",
        "Expose a visible cancel/close button and preserve keyboard accessibility.",
    ])
    add_callout(doc, "Completion definition", "The Business Assistant is production-ready when every request is tenant-scoped, permission-checked, deterministic where possible, resilient without AI, confirmation-gated for writes, and auditable after execution.", GREEN)
    path = OUTPUT_DIR / "Vendly_Business_Assistant_Build_Guide.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for generated_path in [build_chatbot_guide(), build_business_assistant_guide()]:
        print(generated_path)
